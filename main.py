# main.py
import base64
import magic
import io
import pandas as pd
import docx
import docx2txt
from pdf2image import convert_from_bytes
import pdfplumber
import uvicorn
from PIL import Image

from fastapi import FastAPI, HTTPException
from pydantic import BaseModel, Field
from typing import List, Union, Literal, Optional

# --- Pydantic Models for the Unified Structure ---
# Modelos que definem a estrutura de saída unificada para a API.

class ImageData(BaseModel):
    """Define a estrutura para retornar dados de imagem."""
    original_mime_type: str
    image_base64_png: str

class TextBlock(BaseModel):
    """Define um bloco de conteúdo de texto."""
    type: Literal["bloco_texto"] = "bloco_texto"
    source_page: Optional[int] = None
    content: str

class ImageBlock(BaseModel):
    """Define um bloco de conteúdo de imagem."""
    type: Literal["bloco_imagem"] = "bloco_imagem"
    source_page: Optional[int] = None
    content: ImageData

class FileInput(BaseModel):
    """Modelo de entrada para a requisição da API."""
    file_base64: str

class UnifiedProcessResponse(BaseModel):
    """Modelo de resposta unificada para a API."""
    status: Literal["success", "error"]
    content_type: Literal["documento_unificado", "unsupported", "error"]
    data: Union[List[Union[TextBlock, ImageBlock]], None]
    message: str

# --- FastAPI App Initialization ---
app = FastAPI(
    title="Serviço de Extração Unificada de Dados",
    description="Uma API que recebe um arquivo em base64 e retorna uma estrutura de conteúdo unificada para agentes de IA.",
    version="4.0.0",
)

# --- Core Processing Function ---

def process_base64_file(base64_string: str) -> dict:
    """
    Decodifica uma string base64, identifica o tipo de arquivo e retorna
    uma lista unificada de blocos de conteúdo (texto e/ou imagem).
    Usa o mesmo padrão de conversão de PDF para imagem do PDFSplitter (DPI=200, formato PNG).
    """
    try:
        decoded_bytes = base64.b64decode(base64_string)
        mime_type = magic.from_buffer(decoded_bytes, mime=True)
        file_stream = io.BytesIO(decoded_bytes)
        
        content_blocks = []

        # --- PDF Processing (Unified Logic) ---
        if mime_type == 'application/pdf':
            try:
                with pdfplumber.open(file_stream) as pdf:
                    for page in pdf.pages:
                        # Tenta extrair texto da página
                        page_text = page.extract_text(x_tolerance=2)
                        
                        if page_text and page_text.strip():
                            content_blocks.append(TextBlock(source_page=page.page_number, content=page_text.strip()))
                        else:
                            # Se não houver texto, converte a página em imagem PNG seguindo o padrão do PDFSplitter
                            page_image = convert_from_bytes(decoded_bytes, dpi=200, fmt='PNG', first_page=page.page_number, last_page=page.page_number)[0]
                            buffered = io.BytesIO()
                            page_image.save(buffered, format='PNG')
                            buffered.seek(0)
                            img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
                            image_content = ImageData(original_mime_type=mime_type, image_base64_png=img_str)
                            content_blocks.append(ImageBlock(source_page=page.page_number, content=image_content))
                
                return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": f"PDF processado. Total de {len(content_blocks)} páginas."}

            except Exception as pdf_error:
                 raise HTTPException(status_code=500, detail=f"Erro ao processar PDF: {pdf_error}")

        # --- Image Processing ---
        elif mime_type.startswith('image/'):
            img = Image.open(file_stream)
            # Processa a imagem seguindo o padrão do PDFSplitter
            buffered = io.BytesIO()
            # Salva como PNG usando o mesmo formato do PDFSplitter
            img.save(buffered, format='PNG')
            buffered.seek(0)
            png_base64_string = base64.b64encode(buffered.getvalue()).decode('utf-8')
            image_data = ImageData(original_mime_type=mime_type, image_base64_png=png_base64_string)
            content_blocks.append(ImageBlock(content=image_data))
            return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": f"Arquivo de imagem ({mime_type}) processado."}

        # --- DOC Processing (arquivos Word antigos) ---
        elif mime_type == 'application/msword':
            try:
                import olefile
                from oletools.olevba import VBA_Parser
                import re
                
                file_stream.seek(0)
                extracted_text = ""
                
                # Método 1: Tentar extrair usando oletools (mais robusta)
                try:
                    # Verificar se é um arquivo OLE válido
                    if olefile.isOleFile(file_stream):
                        file_stream.seek(0)
                        
                        # Tentar usar oletools para extrair texto
                        vba_parser = VBA_Parser(file_stream)
                        
                        # Se conseguir abrir com oletools, tenta extrair texto simples
                        file_stream.seek(0)
                        ole = olefile.OleFileIO(file_stream)
                        
                        # Listar streams disponíveis
                        streams = ole.listdir()
                        
                        # Procurar por streams que podem conter texto
                        text_content = []
                        for stream_path in streams:
                            try:
                                if any(word in str(stream_path).lower() for word in ['worddocument', '1table', 'data']):
                                    stream_data = ole._olestream(stream_path)
                                    if stream_data:
                                        # Extrair texto legível do stream
                                        readable_text = ''.join([chr(b) if 32 <= b <= 126 else ' ' for b in stream_data])
                                        # Limpar e filtrar texto
                                        words = re.findall(r'\b[a-zA-ZÀ-ÿ0-9.,!?:;\-\s]{3,}\b', readable_text)
                                        if words:
                                            text_content.extend(words[:100])  # Limitar a 100 palavras por stream
                            except:
                                continue
                        
                        ole.close()
                        vba_parser.close()
                        
                        if text_content:
                            extracted_text = ' '.join(text_content[:200])  # Limitar texto total
                            # Limpar texto duplicado e caracteres estranhos
                            extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                            
                except Exception as oletools_error:
                    # Método 2: Extração básica de texto (fallback)
                    try:
                        file_stream.seek(0)
                        raw_data = file_stream.read()
                        
                        # Buscar por padrões de texto em arquivos DOC
                        # Arquivos DOC podem ter texto em codificação específica
                        text_chunks = []
                        
                        # Tentar diferentes encoding
                        for encoding in ['latin1', 'cp1252', 'utf-8', 'utf-16']:
                            try:
                                decoded_text = raw_data.decode(encoding, errors='ignore')
                                # Extrair palavras legíveis
                                words = re.findall(r'\b[a-zA-ZÀ-ÿ0-9.,!?:;\-]{3,}\b', decoded_text)
                                if len(words) > 10:  # Se encontrou palavras suficientes
                                    text_chunks.extend(words[:50])
                                    break
                            except:
                                continue
                        
                        if text_chunks:
                            extracted_text = ' '.join(text_chunks)
                            
                    except Exception as basic_error:
                        extracted_text = ""
                
                # Processar resultado
                if extracted_text and len(extracted_text.strip()) > 20:
                    # Limpeza final do texto
                    extracted_text = re.sub(r'[^\w\s.,!?:;\-àáâãäåçèéêëìíîïñòóôõöùúûüýÿ]', ' ', extracted_text)
                    extracted_text = re.sub(r'\s+', ' ', extracted_text).strip()
                    
                    content_blocks.append(TextBlock(
                        content=f"[TEXTO EXTRAÍDO DE ARQUIVO DOC]\n\n{extracted_text}\n\n"
                               "[NOTA: Extração de arquivo DOC antigo. Para melhor qualidade, converta para DOCX]"
                    ))
                    return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": "Texto extraído de arquivo DOC antigo com sucesso."}
                else:
                    # Se não conseguiu extrair texto suficiente
                    content_blocks.append(TextBlock(
                        content="[ARQUIVO DOC DETECTADO - EXTRAÇÃO LIMITADA]\n\n"
                               "Arquivo Microsoft Word antigo (.doc) detectado, mas não foi possível extrair texto suficiente.\n\n"
                               "Recomendações para melhor extração:\n"
                               "1. Converter para DOCX usando Microsoft Word\n"
                               "2. Salvar como PDF\n"
                               "3. Usar uma ferramenta online de conversão\n\n"
                               f"Formato detectado: {mime_type}\n"
                               f"Tamanho do arquivo: {len(decoded_bytes)} bytes"
                    ))
                    return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": "Arquivo DOC detectado mas extração limitada. Recomenda-se conversão."}
                
            except Exception as doc_error:
                return {"status": "error", "content_type": "error", "data": None, "message": f"Erro ao processar arquivo DOC: {str(doc_error)}"}

        # --- DOCX Processing (arquivos Word modernos) ---
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            try:
                doc = docx.Document(file_stream)
                
                # Lista para armazenar todo o conteúdo extraído
                extracted_content = []
                
                # Extrair texto dos parágrafos
                for para in doc.paragraphs:
                    if para.text.strip():
                        extracted_content.append(para.text.strip())
                
                # Extrair texto das tabelas
                for table in doc.tables:
                    table_content = []
                    for row in table.rows:
                        row_content = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_content.append(cell_text)
                        if row_content:
                            table_content.append(" | ".join(row_content))
                    
                    if table_content:
                        extracted_content.append("\n--- TABELA ---")
                        extracted_content.extend(table_content)
                        extracted_content.append("--- FIM TABELA ---\n")
                
                # Tentar extrair texto de elementos XML diretamente (para capturar conteúdo perdido)
                try:
                    from docx.oxml.text.paragraph import CT_P
                    from docx.oxml.table import CT_Tbl
                    from docx.text.paragraph import Paragraph
                    from docx.table import Table
                    
                    # Percorrer todos os elementos do documento
                    for element in doc.element.body:
                        if isinstance(element, CT_P):
                            # Já capturado pelos parágrafos acima
                            continue
                        elif isinstance(element, CT_Tbl):
                            # Já capturado pelas tabelas acima
                            continue
                        else:
                            # Tentar extrair texto de outros elementos
                            element_text = element.text if hasattr(element, 'text') else ''
                            if element_text and element_text.strip():
                                extracted_content.append(f"[ELEMENTO ESPECIAL]: {element_text.strip()}")
                except Exception as xml_error:
                    # Se falhar a extração XML, continua com o que já foi extraído
                    pass
                
                # Juntar todo o conteúdo
                text_content = '\n'.join(extracted_content)
                
                if text_content and text_content.strip():
                    content_blocks.append(TextBlock(content=text_content.strip()))
                    return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": "Arquivo DOCX processado com extração aprimorada."}
                else:
                    return {"status": "error", "content_type": "error", "data": None, "message": "Documento DOCX está vazio ou não contém texto extraível."}
            except Exception as docx_error:
                return {"status": "error", "content_type": "error", "data": None, "message": f"Erro ao processar arquivo DOCX: {str(docx_error)}"}

        # --- XLSX Processing ---
        elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            df = pd.read_excel(file_stream)
            content_blocks.append(TextBlock(content=df.to_string()))
            return {"status": "success", "content_type": "documento_unificado", "data": content_blocks, "message": "Arquivo XLSX processado."}
        
        # --- Unsupported File Type ---
        else:
            return {"status": "error", "content_type": "unsupported", "data": None, "message": f"Tipo de arquivo não suportado: {mime_type}"}

    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Ocorreu um erro inesperado: {str(e)}")

# --- API Endpoint Definition ---
@app.post("/process-file/", response_model=UnifiedProcessResponse, tags=["File Processing"])
async def create_processing_job(file_input: FileInput):
    """
    Recebe um arquivo codificado em base64 e o transforma em uma estrutura
    de conteúdo unificada, pronta para ser consumida por agentes de IA.
    """
    result = process_base64_file(file_input.file_base64)
    if result["status"] == "error":
        raise HTTPException(status_code=400, detail=result["message"])
    return result

@app.get("/", tags=["Health Check"])
async def read_root():
    """Endpoint raiz para verificar se a API está online."""
    return {"status": "ok", "message": "Serviço de Extração Unificada de Dados está no ar!"}

# --- To run the server directly for testing ---
if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
